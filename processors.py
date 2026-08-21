from __future__ import annotations

import json
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

import fitz
from docx import Document
from docx.text.paragraph import Paragraph

from config import AppConfig
from translator import (
    _exact_translation,
    _reflow_to_source_lines,
    _translate_catalogue,
    _translate_resilient,
    _translate_with_opus,
    _validate_model_translation,
    qa_text,
    should_translate,
    split_long_text,
    translate_text,
)


Progress = Callable[[float, str], None]

_CHECKPOINT_SCHEMA = 4


@dataclass(slots=True)
class PdfTextBlock:
    rect: fitz.Rect
    text: str
    font_size: float
    color: tuple[float, float, float]
    alignment: int
    rotation: int = 0
    expandable: bool = False
    prefer_fast: bool = False
    kind: str = "body"


class TranslationQualityError(RuntimeError):
    """A translation stopped safely and produced a QA report."""

    def __init__(self, message: str, report: Path):
        super().__init__(message)
        self.report = Path(report)


_CJK_SOURCE_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_DRAWING_LABELS = {
    "page": "Стр.",
    "page no": "Стр.",
    "page no.": "Стр.",
}

_SDS_EXACT_TRANSLATIONS = {
    "name of the company": "Наименование компании",
    "e-mail address": "Адрес электронной почты",
    "component": "Компонент",
    "list of carcinogens by the iarc": "Перечень канцерогенов по классификации IARC",
    "stot-single exposure": "STOT — однократное воздействие",
    "may cause drowsiness or dizziness(category 3)": "Может вызывать сонливость или головокружение (категория 3)",
    "stot-repeated exposure": "STOT — многократное воздействие",
    "based on available data, the classification criteria are not met": "На основании имеющихся данных критерии классификации не выполняются",
    "results of pbt and vpvb assessment": "Результаты оценки PBT и vPvB",
    "results of p": "Результаты оценки PBT и vPvB",
    "china inventory of existing chemical substances": "Китайский реестр существующих химических веществ",
    "new zealand inventory of chemicals": "Реестр химических веществ Новой Зеландии",
    "australia. inventory of industrial chemicals (aiic)": "Австралийский реестр промышленных химических веществ (AIIC)",
    "derived no effect level": "Производный безопасный уровень воздействия",
    "lethal dose 50%": "Смертельная доза 50%",
    "effective concentration x%": "Эффективная концентрация X%",
    "very persistent, very bioaccumulative": "Очень стойкие и очень биоаккумулируемые вещества",
}

_SDS_REFERENCE_PREFIXES = {
    "IPCS: The International Chemical Safety Cards (ICSC), website:":
        "IPCS: Международные карты химической безопасности (ICSC), веб-сайт:",
    "IARC, website:": "IARC, веб-сайт:",
    "OECD: The Global Portal to Information on Chemical Substances, website:":
        "OECD: Глобальный портал информации о химических веществах, веб-сайт:",
    "CAMEO Chemicals, website:": "CAMEO Chemicals, веб-сайт:",
    "NLM: ChemIDplus, website:": "NLM: ChemIDplus, веб-сайт:",
    "EPA: Integrated Risk Information System, website:":
        "EPA: Интегрированная информационная система оценки рисков, веб-сайт:",
    "U.S. Department of Transportation: ERG, website:":
        "Министерство транспорта США: ERG, веб-сайт:",
    "Germany GESTIS-database on hazard substance, website:":
        "Немецкая база данных GESTIS по опасным веществам, веб-сайт:",
}

_REPAIR_ACRONYMS = {
    "acgih", "adr", "aiic", "cas", "cfr", "dgr", "dn", "dsl", "ec",
    "echa", "einecs", "encs", "gbz", "ghs", "iata", "iarc", "icao",
    "iecsc", "imdg", "keci", "msds", "niosh", "nzioc", "ntp", "oecd",
    "osha", "pbt", "piccs", "reach", "sds", "stot", "tsca", "un",
    "vpvb",
}


def _normalized_repair_text(text: str) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    return compact.lstrip("|• ").casefold()


def _sds_deterministic_translation(text: str) -> str | None:
    normalized = _normalized_repair_text(text)
    exact = _SDS_EXACT_TRANSLATIONS.get(normalized)
    if exact is not None:
        return exact

    compact = re.sub(r"\s+", " ", text).strip()
    for source_prefix, target_prefix in _SDS_REFERENCE_PREFIXES.items():
        if compact.casefold().startswith(source_prefix.casefold()):
            return target_prefix + compact[len(source_prefix):]

    fire_source = (
        "as in any fire, wear self-contained breathing apparatus"
        "（msha/niosh approved or equivalent) and full protective gear."
    )
    if normalized == fire_source:
        return (
            "Как и при любом пожаре, используйте автономный дыхательный аппарат "
            "(одобренный MSHA/NIOSH или эквивалентный) и полный комплект "
            "защитного снаряжения."
        )
    return None


def _should_repair_line(text: str) -> bool:
    """Recognise source prose without reprocessing Russian acronym rows."""
    stripped = text.strip()
    if not stripped:
        return False
    if _CJK_SOURCE_RE.search(stripped):
        return True
    if re.fullmatch(r"https?://\S+|[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", stripped):
        return False
    if re.search(r"\b(?:CO\.?|LTD\.?|LLC|INC\.?)\s*$", stripped, re.IGNORECASE):
        return False

    latin_words = re.findall(r"\b[A-Za-z][A-Za-z-]{2,}\b", stripped)
    if not latin_words:
        return False
    cyrillic_words = re.findall(r"\b[А-Яа-яЁё][А-Яа-яЁё-]{2,}\b", stripped)
    meaningful = [
        word for word in latin_words
        if word.casefold() not in _REPAIR_ACRONYMS and not word.isupper()
    ]
    if not meaningful:
        return False
    if cyrillic_words:
        return len(meaningful) >= 2 and len(meaningful) > len(cyrillic_words)
    return should_translate(stripped)


def _lines_share_table_row(first: dict, second: dict) -> bool:
    first_rect = fitz.Rect(first["bbox"])
    second_rect = fitz.Rect(second["bbox"])
    overlap = min(first_rect.y1, second_rect.y1) - max(first_rect.y0, second_rect.y0)
    return overlap >= min(first_rect.height, second_rect.height) * 0.55


def _translate_pdf_text(
    text: str,
    config: AppConfig,
    *,
    prefer_fast: bool = False,
) -> str:
    """Quality-first PDF policy with a fast path only for drawing labels."""
    if not should_translate(text):
        return text

    sds_exact = _sds_deterministic_translation(text)
    if sds_exact is not None:
        return sds_exact

    drawing_exact = _DRAWING_LABELS.get(text.strip().casefold())
    if drawing_exact is not None:
        return drawing_exact

    exact = _exact_translation(text)
    if exact is not None:
        return exact
    catalogue = _translate_catalogue(text, config)
    if catalogue is not None:
        return catalogue

    if prefer_fast and not _CJK_SOURCE_RE.search(text):
        try:
            translated = _translate_with_opus(text)
            if translated.strip():
                return translated
        except RuntimeError:
            pass

    translated_chunks: list[str] = []
    for chunk in split_long_text(text):
        translated_chunks.append(
            _translate_resilient(chunk, config)
            if should_translate(chunk)
            else chunk
        )
    translated = _reflow_to_source_lines(text, "".join(translated_chunks))
    _validate_model_translation(text, translated)
    return translated


def _safe_progress(progress: Progress, value: float, message: str) -> None:
    try:
        progress(max(0.0, min(100.0, float(value))), message)
    except Exception:
        # Ошибка интерфейсного callback не должна повреждать перевод файла.
        pass


def _write_qa_report(
    destination: Path,
    warnings: list[str],
    *,
    source: Path | None = None,
    processed_units: int = 0,
) -> Path:
    report = destination.with_name(f"{destination.stem}_QA_REPORT.txt")
    lines = [
        "PDFMathTranslate WLL — отчёт контроля качества",
        f"Исходный файл: {source.name if source else 'не указан'}",
        f"Результат: {destination.name}",
        f"Обработано элементов: {processed_units}",
        "",
    ]

    unique_warnings = list(dict.fromkeys(warnings))
    if unique_warnings:
        lines.append(f"Найдены замечания: {len(unique_warnings)}")
        lines.extend(f"- {warning}" for warning in unique_warnings)
    else:
        lines.append("Автоматическая проверка не выявила явных проблем.")

    report.write_text("\n".join(lines), encoding="utf-8")
    return report


def _paragraphs_in_document(doc: Document) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    seen: set[int] = set()

    def add(paragraph: Paragraph) -> None:
        key = id(paragraph._p)
        if key in seen or not paragraph.text.strip():
            return
        seen.add(key)
        paragraphs.append(paragraph)

    for paragraph in doc.paragraphs:
        add(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add(paragraph)

    # Колонтитулы часто содержат технические заголовки и номера разделов.
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            add(paragraph)
        for paragraph in section.footer.paragraphs:
            add(paragraph)

    return paragraphs


def _replace_paragraph_text(paragraph: Paragraph, translated: str) -> None:
    """
    Сохраняет формат первого текстового run и не разрушает свойства абзаца.
    Для сложного смешанного форматирования это безопаснее, чем paragraph.text.
    """
    if not paragraph.runs:
        paragraph.add_run(translated)
        return

    nonempty_runs = [run for run in paragraph.runs if run.text]
    target = nonempty_runs[0] if nonempty_runs else paragraph.runs[0]
    target.text = translated

    for run in paragraph.runs:
        if run is not target:
            run.text = ""


def translate_docx(
    source: Path,
    destination: Path,
    config: AppConfig,
    progress: Progress,
) -> Path:
    doc = Document(str(source))
    paragraphs = _paragraphs_in_document(doc)

    warnings: list[str] = []
    total = max(len(paragraphs), 1)
    processed = 0

    for index, paragraph in enumerate(paragraphs, start=1):
        original = paragraph.text
        if not should_translate(original):
            _safe_progress(progress, int(index / total * 100), f"DOCX: обработано {index}/{total}")
            continue

        try:
            translated = translate_text(original, config)
            _replace_paragraph_text(paragraph, translated)
            processed += 1

            result = qa_text(
                translated,
                f"DOCX, абзац {index}",
                source_text=original,
            )
            warnings.extend(result.warnings)
        except Exception as exc:
            warnings.append(f"DOCX, абзац {index}: ошибка перевода — {exc}")

        _safe_progress(
            progress,
            int(index / total * 100),
            f"DOCX: обработано {index}/{total}",
        )

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))

    return _write_qa_report(
        destination,
        warnings,
        source=source,
        processed_units=processed,
    )


def _windows_cyrillic_font() -> str | None:
    if os.name != "nt":
        return None

    fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    candidates = (
        "arial.ttf",
        "segoeui.ttf",
        "calibri.ttf",
        "tahoma.ttf",
        "times.ttf",
    )
    for name in candidates:
        candidate = fonts_dir / name
        if candidate.exists():
            return str(candidate)
    return None


def _rgb_from_int(color_value: int | None) -> tuple[float, float, float]:
    if not isinstance(color_value, int):
        return (0.0, 0.0, 0.0)

    red = (color_value >> 16) & 255
    green = (color_value >> 8) & 255
    blue = color_value & 255
    return (red / 255.0, green / 255.0, blue / 255.0)


def _estimate_alignment(block_rect: fitz.Rect, line_rects: list[fitz.Rect]) -> int:
    if not line_rects:
        return fitz.TEXT_ALIGN_LEFT

    avg_left = sum(rect.x0 for rect in line_rects) / len(line_rects)
    avg_right = sum(rect.x1 for rect in line_rects) / len(line_rects)
    left_gap = avg_left - block_rect.x0
    right_gap = block_rect.x1 - avg_right

    if abs(left_gap - right_gap) <= max(2.0, block_rect.width * 0.05):
        return fitz.TEXT_ALIGN_CENTER
    if left_gap > right_gap * 1.8:
        return fitz.TEXT_ALIGN_RIGHT
    return fitz.TEXT_ALIGN_LEFT


def _extract_pdf_blocks(page: fitz.Page) -> list[PdfTextBlock]:
    data = page.get_text(
        "dict",
        flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE,
    )
    result: list[PdfTextBlock] = []

    for block in data.get("blocks", []):
        if block.get("type") != 0 or "bbox" not in block:
            continue

        lines = block.get("lines", [])
        spans = [span for line in lines for span in line.get("spans", [])]
        if not spans:
            continue

        text_lines: list[str] = []
        line_rects: list[fitz.Rect] = []

        for line in lines:
            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
            text_lines.append(line_text.rstrip())
            if "bbox" in line:
                line_rects.append(fitz.Rect(line["bbox"]))

        text = "\n".join(text_lines).strip()
        if not should_translate(text):
            continue

        rect = fitz.Rect(block["bbox"])
        if rect.width < 8 or rect.height < 5:
            continue

        sizes = [float(span.get("size", 10.0)) for span in spans if float(span.get("size", 0)) > 0]
        font_size = sum(sizes) / max(len(sizes), 1)

        colors = [span.get("color") for span in spans if isinstance(span.get("color"), int)]
        color = _rgb_from_int(colors[0] if colors else None)

        direction = lines[0].get("dir", (1.0, 0.0)) if lines else (1.0, 0.0)
        rotation = 0
        if isinstance(direction, (list, tuple)) and len(direction) == 2:
            x, y = float(direction[0]), float(direction[1])
            if abs(y) > abs(x):
                rotation = 90 if y > 0 else 270

        alignment = (
            fitz.TEXT_ALIGN_LEFT
            if len(text_lines) >= 4
            else _estimate_alignment(rect, line_rects)
        )

        result.append(
            PdfTextBlock(
                rect=rect,
                text=text,
                font_size=font_size,
                color=color,
                alignment=alignment,
                rotation=rotation,
            )
        )

    return result


def _line_to_pdf_block(
    line: dict,
    *,
    rect: fitz.Rect | None = None,
    text: str | None = None,
    expandable: bool = False,
    prefer_fast: bool = False,
    kind: str = "body",
) -> PdfTextBlock | None:
    spans = line.get("spans", [])
    if not spans or "bbox" not in line:
        return None
    line_text = text if text is not None else "".join(
        span.get("text", "") for span in spans
    ).strip()
    if not should_translate(line_text):
        return None
    line_rect = fitz.Rect(rect or line["bbox"])
    if line_rect.width < 4 or line_rect.height < 3:
        return None
    sizes = [
        float(span.get("size", 10.0))
        for span in spans
        if float(span.get("size", 0.0)) > 0
    ]
    colors = [
        span.get("color")
        for span in spans
        if isinstance(span.get("color"), int)
    ]
    direction = line.get("dir", (1.0, 0.0))
    rotation = 0
    if isinstance(direction, (list, tuple)) and len(direction) == 2:
        x, y = float(direction[0]), float(direction[1])
        if abs(y) > abs(x):
            rotation = 90 if y > 0 else 270
    return PdfTextBlock(
        rect=line_rect,
        text=line_text,
        font_size=sum(sizes) / max(1, len(sizes)),
        color=_rgb_from_int(colors[0] if colors else None),
        alignment=fitz.TEXT_ALIGN_LEFT,
        rotation=rotation,
        expandable=expandable,
        prefer_fast=prefer_fast,
        kind=kind,
    )


def _extract_schematic_repair_blocks(data: dict) -> list[PdfTextBlock]:
    """Collapse bilingual drawing labels to one English→Russian operation."""
    result: list[PdfTextBlock] = []
    for raw_block in data.get("blocks", []):
        if raw_block.get("type") != 0:
            continue
        lines = [
            line for line in raw_block.get("lines", [])
            if line.get("spans") and "bbox" in line
        ]
        index = 0
        while index < len(lines):
            line = lines[index]
            current = "".join(
                span.get("text", "") for span in line.get("spans", [])
            ).strip()
            current_rect = fitz.Rect(line["bbox"])

            if re.search(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]", current):
                # Many CAD exports put both languages in one text object
                # ("机械手联锁 Safety Inter Locking").  Translate the English
                # duplicate and redact the complete bilingual object.  This
                # is much faster and more reliable than asking the model to
                # interpret hundreds of isolated Chinese labels.
                latin_candidates = [
                    match.group(0).strip()
                    for match in re.finditer(
                        r"[A-Za-z][A-Za-z0-9 .,/()_:+-]*",
                        current,
                    )
                ]
                latin_duplicate = next(
                    (
                        candidate
                        for candidate in latin_candidates
                        if should_translate(candidate)
                    ),
                    "",
                )
                if should_translate(latin_duplicate):
                    block = _line_to_pdf_block(
                        line,
                        rect=current_rect,
                        text=latin_duplicate,
                        expandable=True,
                        prefer_fast=True,
                        kind="schematic",
                    )
                    if block is not None:
                        result.append(block)
                    index += 1
                    continue

                if index + 1 < len(lines):
                    following = lines[index + 1]
                    following_text = "".join(
                        span.get("text", "")
                        for span in following.get("spans", [])
                    ).strip()
                    following_rect = fitz.Rect(following["bbox"])
                    same_label = (
                        should_translate(following_text)
                        and not re.search(
                            r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]",
                            following_text,
                        )
                        and abs(current_rect.x0 - following_rect.x0)
                        <= max(20.0, current_rect.height * 1.5)
                        and following_rect.y0 - current_rect.y1
                        <= max(2.0, current_rect.height * 0.45)
                    )
                    if same_label:
                        pair_rect = current_rect | following_rect
                        block = _line_to_pdf_block(
                            following,
                            rect=pair_rect,
                            text=following_text,
                            expandable=True,
                            prefer_fast=True,
                            kind="schematic",
                        )
                        if block is not None:
                            result.append(block)
                        index += 2
                        continue

                block = _line_to_pdf_block(
                    line,
                    expandable=True,
                    kind="schematic_cjk",
                )
            else:
                block = _line_to_pdf_block(
                    line,
                    expandable=True,
                    prefer_fast=True,
                    kind="schematic",
                )
            if block is not None:
                result.append(block)
            index += 1
    return result


def _extract_pdf_repair_blocks(page: fitz.Page) -> list[PdfTextBlock]:
    """Extract residual text as paragraphs/cells or bilingual drawing labels."""
    data = page.get_text(
        "dict",
        flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE,
    )
    all_lines = [
        line
        for raw_block in data.get("blocks", [])
        if raw_block.get("type") == 0
        for line in raw_block.get("lines", [])
        if line.get("spans") and "bbox" in line
    ]
    cjk_lines = sum(
        bool(re.search(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]", "".join(
            span.get("text", "") for span in line.get("spans", [])
        )))
        for line in all_lines
    )
    if cjk_lines >= 2 and len(all_lines) >= 20:
        return _extract_schematic_repair_blocks(data)

    result: list[PdfTextBlock] = []

    for raw_block in data.get("blocks", []):
        if raw_block.get("type") != 0:
            continue
        lines = raw_block.get("lines", [])
        group: list[dict] = []

        def flush_group() -> None:
            if not group:
                return
            text_lines = [
                "".join(span.get("text", "") for span in line.get("spans", [])).strip()
                for line in group
            ]
            text = "\n".join(text_lines).strip()
            rects = [fitz.Rect(line["bbox"]) for line in group]
            rect = fitz.Rect(rects[0])
            for item in rects[1:]:
                rect |= item
            block = _line_to_pdf_block(
                group[0],
                rect=rect,
                text=text,
                expandable=len(group) == 1,
                kind="paragraph" if len(group) > 1 else "cell",
            )
            if block is not None:
                result.append(block)
            group.clear()

        for line in lines:
            spans = line.get("spans", [])
            text = "".join(span.get("text", "") for span in spans).strip()
            if spans and "bbox" in line and _should_repair_line(text):
                if group and _lines_share_table_row(group[-1], line):
                    flush_group()
                group.append(line)
            else:
                flush_group()
        flush_group()
    return result


def _expanded_rect(rect: fitz.Rect, page_rect: fitz.Rect, margin: float = 0.6) -> fitz.Rect:
    expanded = fitz.Rect(
        max(page_rect.x0, rect.x0 - margin),
        max(page_rect.y0, rect.y0 - margin),
        min(page_rect.x1, rect.x1 + margin),
        min(page_rect.y1, rect.y1 + margin),
    )
    return expanded


def _insert_translated_text(
    page: fitz.Page,
    block: PdfTextBlock,
    text: str,
) -> tuple[bool, float]:
    fontfile = _windows_cyrillic_font()
    fontname = "wllfont" if fontfile else "helv"

    start_size = min(max(block.font_size, 6.0), 22.0)
    minimum_size = 3.2 if block.kind.startswith("schematic") else 3.5
    size = start_size

    rect = _expanded_rect(block.rect, page.rect, margin=0.5)

    while size >= minimum_size:
        result = page.insert_textbox(
            rect,
            text,
            fontsize=size,
            fontname=fontname,
            fontfile=fontfile,
            align=block.alignment,
            color=block.color,
            lineheight=1.0,
            rotate=block.rotation,
            overlay=True,
        )
        if result >= 0:
            return True, size
        size -= 0.35

    # Последняя попытка в слегка увеличенной области, не выходящей за страницу.
    fallback_rect = _expanded_rect(block.rect, page.rect, margin=2.5)
    if block.expandable and block.rotation == 0:
        if block.kind.startswith("schematic"):
            # Drawing labels have free horizontal space around the source.
            fallback_rect.x1 = min(
                page.rect.x1 - 1.0,
                max(
                    fallback_rect.x1,
                    block.rect.x0 + 120.0,
                    block.rect.x1 + block.rect.width,
                ),
            )
            fallback_rect.y1 = min(
                page.rect.y1 - 1.0,
                max(fallback_rect.y1, block.rect.y1 + block.rect.height * 0.8),
            )
        else:
            # A table cell may use the whitespace after the extracted glyphs,
            # but expansion is capped so it cannot reach the next column.
            extra_width = min(100.0, max(20.0, block.rect.width * 0.65))
            fallback_rect.x1 = min(
                page.rect.x1 - 1.0,
                max(fallback_rect.x1, block.rect.x1 + extra_width),
            )
    result = page.insert_textbox(
        fallback_rect,
        text,
        fontsize=minimum_size,
        fontname=fontname,
        fontfile=fontfile,
        align=block.alignment,
        color=block.color,
        lineheight=0.88,
        rotate=block.rotation,
        overlay=True,
    )
    return result >= 0, minimum_size


def _translation_fits(page_rect: fitz.Rect, block: PdfTextBlock, text: str) -> bool:
    """Check layout on a disposable page before the source block is redacted."""
    probe = fitz.open()
    try:
        page = probe.new_page(width=page_rect.width, height=page_rect.height)
        inserted, _ = _insert_translated_text(page, block, text)
        return inserted
    except Exception:
        return False
    finally:
        probe.close()


def _page_has_images(page: fitz.Page) -> bool:
    try:
        return bool(page.get_images(full=True))
    except Exception:
        return False


def _format_duration(seconds: float) -> str:
    seconds = max(0, int(seconds))
    hours, remainder = divmod(seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours} ч {minutes:02d} мин"
    if minutes:
        return f"{minutes} мин {secs:02d} сек"
    return f"{secs} сек"


def _progress_message(
    base: str,
    *,
    started: float,
    session_fraction: float,
) -> str:
    elapsed = time.monotonic() - started
    message = f"{base} · прошло {_format_duration(elapsed)}"
    if session_fraction > 0.0:
        remaining = elapsed * max(0.0, 1.0 - session_fraction) / session_fraction
        message += f" · осталось примерно {_format_duration(remaining)}"
    return message


def _checkpoint_paths(destination: Path) -> tuple[Path, Path]:
    prefix = f".{destination.name}.wll"
    return (
        destination.with_name(prefix + "-part.pdf"),
        destination.with_name(prefix + "-state.json"),
    )


def _source_signature(source: Path) -> dict[str, object]:
    stat = source.stat()
    return {
        "path": str(source.resolve()),
        "size": stat.st_size,
        "mtime_ns": stat.st_mtime_ns,
    }


def _discard_checkpoint(partial: Path, state_path: Path) -> None:
    partial.unlink(missing_ok=True)
    state_path.unlink(missing_ok=True)


def _load_checkpoint(
    source: Path,
    partial: Path,
    state_path: Path,
    *,
    start_index: int,
    end_index: int,
    repair_mode: bool = False,
) -> tuple[fitz.Document, int, list[str], int, bool]:
    if not partial.exists() or not state_path.exists():
        return fitz.open(str(source)), start_index, [], 0, False

    try:
        state = json.loads(state_path.read_text(encoding="utf-8"))
        next_index = int(state["next_page_index"])
        valid = (
            state.get("schema") == _CHECKPOINT_SCHEMA
            and state.get("source") == _source_signature(source)
            and int(state.get("start_index", -1)) == start_index
            and int(state.get("end_index", -1)) == end_index
            and bool(state.get("repair_mode", False)) == repair_mode
            and start_index <= next_index <= end_index + 1
        )
        if not valid:
            raise ValueError("контрольные данные не соответствуют текущему заданию")

        # Открываем из памяти, чтобы Windows не удерживала checkpoint-файл.
        doc = fitz.open(stream=partial.read_bytes(), filetype="pdf")
        warnings = [str(item) for item in state.get("warnings", [])]
        processed_blocks = max(0, int(state.get("processed_blocks", 0)))
        return doc, next_index, warnings, processed_blocks, True
    except Exception:
        _discard_checkpoint(partial, state_path)
        return fitz.open(str(source)), start_index, [], 0, False


def _save_checkpoint(
    doc: fitz.Document,
    source: Path,
    partial: Path,
    state_path: Path,
    *,
    start_index: int,
    end_index: int,
    next_page_index: int,
    warnings: list[str],
    processed_blocks: int,
    repair_mode: bool = False,
) -> fitz.Document:
    partial.parent.mkdir(parents=True, exist_ok=True)
    partial_tmp = partial.with_name(partial.stem + "-tmp.pdf")
    state_tmp = state_path.with_name(state_path.name + ".tmp")
    partial_tmp.unlink(missing_ok=True)
    state_tmp.unlink(missing_ok=True)

    reopened: fitz.Document | None = None
    try:
        doc.save(str(partial_tmp), garbage=3, deflate=True, clean=False)

        # garbage=3 reorganises the PDF xref table.  Continuing to edit the
        # old in-memory document can leave page resources pointing at the old
        # object numbers ("object out of range" on the next page).  Validate
        # and reopen the exact checkpoint bytes before processing continues.
        reopened = fitz.open(stream=partial_tmp.read_bytes(), filetype="pdf")
        if reopened.page_count != doc.page_count:
            raise RuntimeError("контрольная копия PDF содержит неверное число страниц")

        os.replace(partial_tmp, partial)

        payload = {
            "schema": _CHECKPOINT_SCHEMA,
            "source": _source_signature(source),
            "start_index": start_index,
            "end_index": end_index,
            "next_page_index": next_page_index,
            "warnings": warnings,
            "processed_blocks": processed_blocks,
            "repair_mode": repair_mode,
        }
        state_tmp.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        os.replace(state_tmp, state_path)
        return reopened
    except Exception:
        if reopened is not None:
            reopened.close()
        partial_tmp.unlink(missing_ok=True)
        state_tmp.unlink(missing_ok=True)
        raise


def _continue_from_checkpoint(doc: fitz.Document, checkpoint: fitz.Document) -> fitz.Document:
    """Close the pre-save document only after its checkpoint was validated."""
    doc.close()
    return checkpoint


def _save_final_pdf(doc: fitz.Document, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.wll-final-tmp.pdf")
    temporary.unlink(missing_ok=True)
    doc.save(
        str(temporary),
        garbage=4,
        deflate=True,
        clean=True,
        pretty=False,
    )
    os.replace(temporary, destination)


def translate_pdf(
    source: Path,
    destination: Path,
    config: AppConfig,
    progress: Progress,
    *,
    page_start: int | None = None,
    page_end: int | None = None,
    repair_mode: bool = False,
) -> Path:
    source = Path(source)
    destination = Path(destination)

    with fitz.open(str(source)) as source_doc:
        total_pages = source_doc.page_count

    if total_pages < 1:
        raise ValueError("PDF не содержит страниц")

    start_index = max(0, int(page_start or 1) - 1)
    end_index = min(total_pages - 1, int(page_end or total_pages) - 1)
    if start_index > end_index:
        raise ValueError("Начальная страница диапазона больше конечной")

    partial, state_path = _checkpoint_paths(destination)
    doc, resume_index, warnings, processed_blocks, resumed = _load_checkpoint(
        source,
        partial,
        state_path,
        start_index=start_index,
        end_index=end_index,
        repair_mode=repair_mode,
    )
    range_pages = end_index - start_index + 1
    session_pages = max(1, end_index - resume_index + 1)
    started = time.monotonic()
    translation_cache: dict[tuple[str, bool], str] = {}

    if resumed:
        _safe_progress(
            progress,
            ((resume_index - start_index) / range_pages) * 100.0,
            f"Продолжение с сохранённой страницы {resume_index + 1}",
        )

    try:
        for page_index in range(resume_index, end_index + 1):
            page = doc[page_index]
            blocks = (
                _extract_pdf_repair_blocks(page)
                if repair_mode
                else _extract_pdf_blocks(page)
            )
            page_offset = page_index - start_index
            session_offset = page_index - resume_index

            if not blocks:
                if repair_mode:
                    _safe_progress(
                        progress,
                        ((page_offset + 1) / range_pages) * 100.0,
                        _progress_message(
                            f"PDF: страница {page_index + 1}/{doc.page_count} уже переведена",
                            started=started,
                            session_fraction=(session_offset + 1) / session_pages,
                        ),
                    )
                    doc = _continue_from_checkpoint(
                        doc,
                        _save_checkpoint(
                            doc,
                            source,
                            partial,
                            state_path,
                            start_index=start_index,
                            end_index=end_index,
                            next_page_index=page_index + 1,
                            warnings=warnings,
                            processed_blocks=processed_blocks,
                            repair_mode=repair_mode,
                        ),
                    )
                    continue
                image_note = " На странице есть изображения." if _page_has_images(page) else ""
                warnings.append(
                    f"PDF, страница {page_index + 1}: текстовый слой не найден; "
                    f"возможно, требуется OCR.{image_note}"
                )
                _safe_progress(
                    progress,
                    ((page_offset + 1) / range_pages) * 100.0,
                    _progress_message(
                        f"PDF: завершена страница {page_index + 1}/{doc.page_count}",
                        started=started,
                        session_fraction=(session_offset + 1) / session_pages,
                    ),
                )
                doc = _continue_from_checkpoint(
                    doc,
                    _save_checkpoint(
                        doc,
                        source,
                        partial,
                        state_path,
                        start_index=start_index,
                        end_index=end_index,
                        next_page_index=page_index + 1,
                        warnings=warnings,
                        processed_blocks=processed_blocks,
                        repair_mode=repair_mode,
                    ),
                )
                continue

            # Ошибка одного блока не должна останавливать большой документ.
            # Неудачный блок остаётся в оригинале и фиксируется в QA-отчёте.
            translated_blocks: list[tuple[int, PdfTextBlock, str]] = []
            for block_number, block in enumerate(blocks, start=1):
                try:
                    cache_key = (block.text, block.prefer_fast)
                    if cache_key in translation_cache:
                        translated = translation_cache[cache_key]
                    else:
                        translated = _translate_pdf_text(
                            block.text,
                            config,
                            prefer_fast=block.prefer_fast,
                        )
                        translation_cache[cache_key] = translated
                    translated_blocks.append((block_number, block, translated))
                    warnings.extend(
                        qa_text(
                            translated,
                            f"PDF, страница {page_index + 1}, блок {block_number}",
                            source_text=block.text,
                        ).warnings
                    )
                except Exception as exc:
                    warnings.append(
                        f"PDF, страница {page_index + 1}, блок {block_number}: "
                        f"ошибка перевода — {exc}; исходный блок оставлен без "
                        "изменения, обработка документа продолжена"
                    )

                _safe_progress(
                    progress,
                    ((page_offset + block_number / len(blocks)) / range_pages) * 100.0,
                    _progress_message(
                        f"PDF: страница {page_index + 1}/{doc.page_count}, "
                        f"блок {block_number}/{len(blocks)}",
                        started=started,
                        session_fraction=(
                            session_offset + block_number / len(blocks)
                        ) / session_pages,
                    ),
                )

            if translated_blocks:
                insertable_blocks: list[tuple[int, PdfTextBlock, str]] = []
                for block_number, block, translated in translated_blocks:
                    if _translation_fits(page.rect, block, translated):
                        insertable_blocks.append((block_number, block, translated))
                    else:
                        warnings.append(
                            f"PDF, страница {page_index + 1}, блок {block_number}: "
                            "русский текст не помещается в исходную область; "
                            "исходный блок сохранён, обработка документа продолжена"
                        )

                # Удаляем только успешно переведённые блоки.
                for _, block, _ in insertable_blocks:
                    page.add_redact_annot(
                        _expanded_rect(block.rect, page.rect, margin=0.4),
                        fill=(1, 1, 1),
                    )

                if insertable_blocks:
                    original_page = fitz.open()
                    original_page.insert_pdf(
                        doc,
                        from_page=page_index,
                        to_page=page_index,
                    )
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                    inserted_on_page = 0
                    try:
                        for block_number, block, translated in insertable_blocks:
                            inserted, used_size = _insert_translated_text(
                                page, block, translated
                            )

                            if not inserted:
                                # Redaction is destructive.  Restore the whole
                                # original page so a rare renderer mismatch can
                                # never leave a blank area in the final PDF.
                                doc.delete_page(page_index)
                                doc.insert_pdf(original_page, start_at=page_index)
                                warnings.append(
                                    f"PDF, страница {page_index + 1}, "
                                    f"блок {block_number}: неожиданная ошибка "
                                    "вставки после проверки макета; исходная "
                                    "страница восстановлена, обработка документа "
                                    "продолжена"
                                )
                                inserted_on_page = 0
                                break

                            inserted_on_page += 1
                            if used_size < max(5.0, block.font_size * 0.55):
                                warnings.append(
                                    f"PDF, страница {page_index + 1}, "
                                    f"блок {block_number}: шрифт сильно уменьшен "
                                    f"до {used_size:.1f} pt"
                                )
                    finally:
                        original_page.close()

                    processed_blocks += inserted_on_page

            _safe_progress(
                progress,
                ((page_offset + 1) / range_pages) * 100.0,
                _progress_message(
                    f"PDF: завершена страница {page_index + 1}/{doc.page_count}; "
                    "результат страницы сохранён",
                    started=started,
                    session_fraction=(session_offset + 1) / session_pages,
                ),
            )
            doc = _continue_from_checkpoint(
                doc,
                _save_checkpoint(
                    doc,
                    source,
                    partial,
                    state_path,
                    start_index=start_index,
                    end_index=end_index,
                    next_page_index=page_index + 1,
                    warnings=warnings,
                    processed_blocks=processed_blocks,
                    repair_mode=repair_mode,
                ),
            )

        _save_final_pdf(doc, destination)
        _discard_checkpoint(partial, state_path)
    finally:
        doc.close()

    return _write_qa_report(
        destination,
        warnings,
        source=source,
        processed_units=processed_blocks,
    )
